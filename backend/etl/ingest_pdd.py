"""
Ingest PDD Export (.docx) -> pdd_variants.json (cached)

Extracts paragraph-level sections from the Process Definition Document.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import docx

from .config import CACHE_DIR, DATASET_DIR, ensure_dirs


def ingest_pdd(force: bool = False) -> Path:
    ensure_dirs()
    out = CACHE_DIR / "pdd_variants.json"
    if not force and out.exists():
        return out

    docx_path = None
    for f in sorted(DATASET_DIR.iterdir()):
        if f.suffix == ".docx" and "PDD" in f.name:
            docx_path = f
            break

    if docx_path is None:
        raise FileNotFoundError("PDD .docx not found in dataset")

    doc = docx.Document(str(docx_path))

    sections: list[dict] = []
    current_heading = ""
    current_body: list[str] = []

    for para in doc.paragraphs:
        # Detect headings
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Flush previous section
            if current_heading or current_body:
                sections.append({
                    "heading": current_heading,
                    "body": "\n".join(current_body).strip(),
                })
            current_heading = para.text.strip()
            current_body = []
        else:
            text = para.text.strip()
            if text:
                current_body.append(text)

    # Flush last section
    if current_heading or current_body:
        sections.append({
            "heading": current_heading,
            "body": "\n".join(current_body).strip(),
        })

    # Also extract tables for structured data
    tables_data: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables_data.append(rows)

    result = {
        "sections": sections,
        "tables": tables_data,
        "source_file": docx_path.name,
    }

    with open(out, "w", encoding="utf-8") as fh:
        json.dump(result, fh, indent=2, ensure_ascii=False)

    print(f"[ingest_pdd] {len(sections)} sections, {len(tables_data)} tables -> {out.name}")
    return out


if __name__ == "__main__":
    ingest_pdd(force="--force" in sys.argv)
